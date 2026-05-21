#!/usr/bin/env python3
"""
Генератор статьи для Политехнического молодёжного журнала МГТУ (ptsj.bmstu.ru)
v4: стиль по образцу реальной статьи из журнала (846.pdf)
"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '.pylibs'))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

FONT = 'Times New Roman'
SIZE = Pt(12)
INDENT = Cm(0.7)
NBSP = '\u00a0'
EMDASH = '\u2014'
LAQUO = '\u00ab'
RAQUO = '\u00bb'

doc = Document()

# ── Секция ──
sec = doc.sections[0]
sec.orientation = WD_ORIENT.PORTRAIT
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2)
sec.top_margin = Cm(2)
sec.right_margin = Cm(1.5)
sec.bottom_margin = Cm(1.5)

# ── Стиль ──
sty = doc.styles['Normal']
sty.font.name = FONT
sty.font.size = SIZE
sty.font.color.rgb = RGBColor(0, 0, 0)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.space_after = Pt(0)
sty.paragraph_format.space_before = Pt(0)
sty.paragraph_format.first_line_indent = INDENT

rpr = sty.element.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
    rf.set(qn(a), FONT)
rpr.append(rf)


def SF(run):
    """Set font on run for Cyrillic."""
    r = run._element
    rp = r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    rp.insert(0, rf)


def add_page_numbers():
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._element.append(fld)


def P(text, bold=False, italic=False, align=None, indent=True,
      sz=None, after=None, before=None):
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if align:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = sz or SIZE
    run.bold = bold
    run.italic = italic
    SF(run)
    return p


def RP(parts, align=None, indent=True, after=None, before=None):
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if align:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = sz if (sz := None) else SIZE
        run.bold = bold
        run.italic = italic
        SF(run)
    return p


def SEC(title, text):
    """Раздел в стиле журнала: жирный заголовок + текст на той же строке."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    # Заголовок жирным
    r1 = p.add_run(f'{title}. ')
    r1.font.name = FONT
    r1.font.size = SIZE
    r1.bold = True
    SF(r1)
    # Текст обычным
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = SIZE
    SF(r2)
    return p


def OMML(formula_text, number=None):
    """OMML formula."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    oMathPara = etree.SubElement(p._element, qn('m:oMathPara'))
    oMath = etree.SubElement(oMathPara, qn('m:oMath'))
    r = etree.SubElement(oMath, qn('m:r'))
    rPr = etree.SubElement(r, qn('m:rPr'))
    sty_el = etree.SubElement(rPr, qn('m:sty'))
    sty_el.set(qn('m:val'), 'p')
    w_rPr = etree.SubElement(r, qn('w:rPr'))
    w_rf = etree.SubElement(w_rPr, qn('w:rFonts'))
    w_rf.set(qn('w:ascii'), 'Cambria Math')
    w_rf.set(qn('w:hAnsi'), 'Cambria Math')
    t = etree.SubElement(r, qn('m:t'))
    t.text = formula_text

    if number:
        tab_run = p.add_run(f'\t({number})')
        tab_run.font.name = FONT
        tab_run.font.size = SIZE
        SF(tab_run)
    return p


def remove_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def abstract_block(abstract_text, keywords_text, received_date, copyright_text):
    """Двухколоночная таблица: аннотация слева, ключевые слова справа."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_borders(table)

    # Ширина колонок (примерно 65% / 35%)
    for cell in table.columns[0].cells:
        cell.width = Cm(11)
    for cell in table.columns[1].cells:
        cell.width = Cm(6.5)

    # Левая колонка — аннотация
    left = table.cell(0, 0)
    left.text = ''
    lp = left.paragraphs[0]
    lp.paragraph_format.first_line_indent = Cm(0)
    lp.paragraph_format.space_after = Pt(0)
    r1 = lp.add_run('Аннотация')
    r1.font.name = FONT
    r1.font.size = SIZE
    r1.bold = True
    SF(r1)

    lp2 = left.add_paragraph()
    lp2.paragraph_format.first_line_indent = Cm(0)
    lp2.paragraph_format.space_after = Pt(0)
    r2 = lp2.add_run(abstract_text)
    r2.font.name = FONT
    r2.font.size = SIZE
    SF(r2)

    # Правая колонка — ключевые слова
    right = table.cell(0, 1)
    right.text = ''
    rp = right.paragraphs[0]
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.space_after = Pt(4)
    r3 = rp.add_run('Ключевые слова')
    r3.font.name = FONT
    r3.font.size = SIZE
    r3.bold = True
    SF(r3)

    rp2 = right.add_paragraph()
    rp2.paragraph_format.first_line_indent = Cm(0)
    rp2.paragraph_format.space_after = Pt(8)
    r4 = rp2.add_run(keywords_text)
    r4.font.name = FONT
    r4.font.size = SIZE
    SF(r4)

    # Дата поступления
    rp3 = right.add_paragraph()
    rp3.paragraph_format.first_line_indent = Cm(0)
    rp3.paragraph_format.space_after = Pt(2)
    r5 = rp3.add_run(f'Поступила в редакцию {received_date}')
    r5.font.name = FONT
    r5.font.size = Pt(10)
    r5.italic = True
    SF(r5)

    # Copyright
    rp4 = right.add_paragraph()
    rp4.paragraph_format.first_line_indent = Cm(0)
    r6 = rp4.add_run(copyright_text)
    r6.font.name = FONT
    r6.font.size = Pt(10)
    SF(r6)

    return table


def abstract_block_en(abstract_text, keywords_text, received_date, copyright_text):
    """English version of abstract block."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_borders(table)

    for cell in table.columns[0].cells:
        cell.width = Cm(11)
    for cell in table.columns[1].cells:
        cell.width = Cm(6.5)

    left = table.cell(0, 0)
    left.text = ''
    lp = left.paragraphs[0]
    lp.paragraph_format.first_line_indent = Cm(0)
    lp.paragraph_format.space_after = Pt(0)
    r1 = lp.add_run('Abstract')
    r1.font.name = FONT
    r1.font.size = SIZE
    r1.bold = True
    SF(r1)

    lp2 = left.add_paragraph()
    lp2.paragraph_format.first_line_indent = Cm(0)
    r2 = lp2.add_run(abstract_text)
    r2.font.name = FONT
    r2.font.size = SIZE
    SF(r2)

    right = table.cell(0, 1)
    right.text = ''
    rp = right.paragraphs[0]
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.space_after = Pt(4)
    r3 = rp.add_run('Keywords')
    r3.font.name = FONT
    r3.font.size = SIZE
    r3.bold = True
    SF(r3)

    rp2 = right.add_paragraph()
    rp2.paragraph_format.first_line_indent = Cm(0)
    rp2.paragraph_format.space_after = Pt(8)
    r4 = rp2.add_run(keywords_text)
    r4.font.name = FONT
    r4.font.size = SIZE
    SF(r4)

    rp3 = right.add_paragraph()
    rp3.paragraph_format.first_line_indent = Cm(0)
    rp3.paragraph_format.space_after = Pt(2)
    r5 = rp3.add_run(f'Received {received_date}')
    r5.font.name = FONT
    r5.font.size = Pt(10)
    r5.italic = True
    SF(r5)

    rp4 = right.add_paragraph()
    rp4.paragraph_format.first_line_indent = Cm(0)
    r6 = rp4.add_run(copyright_text)
    r6.font.name = FONT
    r6.font.size = Pt(10)
    SF(r6)

    return table


def make_table(headers, rows, caption=None):
    """Таблица с заголовком НАД ней."""
    if caption:
        P(caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          indent=False, before=6, after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        cp = cell.paragraphs[0]
        cp.paragraph_format.first_line_indent = Cm(0)
        cp.paragraph_format.space_after = Pt(0)
        run = cp.add_run(h)
        run.font.name = FONT
        run.font.size = Pt(10)
        run.bold = True
        SF(run)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.paragraph_format.first_line_indent = Cm(0)
            cp.paragraph_format.space_after = Pt(0)
            run = cp.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(10)
            SF(run)
    return table


# ══════════════════════════════════════════════════════════
#  РУССКИЙ БЛОК
# ══════════════════════════════════════════════════════════

# ── Шапка ──
P('УДК 004.414.23 + 621.865.8', indent=False, after=2)

P('ФОРМАЛЬНАЯ ВЕРИФИКАЦИЯ КООРДИНАЦИИ ПОХОДКИ\n'
  'ШЕСТИНОГОГО РОБОТА МЕТОДОМ СЕТЕЙ ПЕТРИ',
  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=4)

P(f'В.П.{NBSP}Романов', indent=False, after=0)
P('romanov.vp@student.bmstu.ru', indent=False, after=2, sz=Pt(11))
P(f'МГТУ им.{NBSP}Н.Э.{NBSP}Баумана, Москва, Российская Федерация',
  indent=False, after=8)

# ── Аннотация (двухколоночная) ──
abstract_ru = (
    f'Рассматривается задача формальной верификации алгоритма триподной '
    f'походки шестиногого шагающего робота (гексапода) с использованием '
    f'аппарата сетей Петри. Шесть ног робота разделены на две группы '
    f'по три ноги, поочерёдно выполняющие фазы переноса и опоры. '
    f'Построена сеть Петри, моделирующая координацию групп ног '
    f'с ограничением статической устойчивости: в каждый момент времени '
    f'не менее трёх ног должны находиться в контакте с поверхностью. '
    f'Методом S-инвариантов формально доказаны взаимное исключение '
    f'фаз переноса и отсутствие тупиковых ситуаций. Выполнен '
    f'темпоральный анализ с детерминированными задержками: определены '
    f'время цикла походки, максимальная скорость перемещения и узкое '
    f'место. Показано, что применение быстрых сервоприводов сокращает '
    f'период походки на 17{NBSP}% и увеличивает скорость робота '
    f'со 200 до 240{NBSP}мм/с.'
)

keywords_ru = (
    'Сети Петри, шагающий робот, гексапод, триподная походка, '
    'формальная верификация, S-инварианты, координация походки, '
    'статическая устойчивость, темпоральный анализ, взаимное исключение'
)

abstract_block(abstract_ru, keywords_ru,
               '21.05.2026',
               f'\u00a9 МГТУ им.{NBSP}Н.Э.{NBSP}Баумана, 2026')

P('', after=6)

# ══════════════ ВВЕДЕНИЕ ══════════════

SEC('Введение',
    f'Шестиногие шагающие роботы (гексаподы) обладают высокой проходимостью '
    f'и статической устойчивостью, что делает их перспективными для '
    f'работы в неструктурированных средах: пересечённая местность, '
    f'аварийные здания, трубопроводы [1]. Ключевой задачей при разработке '
    f'гексапода является координация шести ног{NBSP}{EMDASH} алгоритм походки, '
    f'который определяет, какие ноги в какой момент отрываются от '
    f'поверхности и переносятся вперёд.')

P(f'Наиболее распространённая походка гексапода{NBSP}{EMDASH} триподная '
  f'(tripod gait): шесть ног разделены на две группы по три (1-3-5 '
  f'и 2-4-6), которые поочерёдно выполняют фазы переноса и опоры [2]. '
  f'Триподная походка обеспечивает максимальную скорость при '
  f'гарантированной статической устойчивости: три опорные ноги всегда '
  f'образуют треугольник, внутри которого находится проекция центра '
  f'масс робота.')

P(f'Однако при программной реализации координация шести параллельно '
  f'работающих приводов порождает класс ошибок, характерных для '
  f'параллельных систем. Гонка состояний (race condition) может '
  f'привести к одновременному отрыву обеих групп ног, что вызовет '
  f'потерю устойчивости и падение робота. Тупиковая ситуация (deadlock) '
  f'может заблокировать движение, если обе группы бесконечно ожидают '
  f'друг друга [3]. Такие ошибки сложно обнаружить тестированием, '
  f'поскольку они возникают лишь при определённых временны\u0301х '
  f'соотношениях между процессами.')

P(f'Конечные автоматы (FSM), традиционно используемые для описания '
  f'логики управления, плохо приспособлены к моделированию параллелизма: '
  f'два процесса с N и M состояниями порождают N{NBSP}\u00d7{NBSP}M '
  f'комбинаций. Для шести ног даже при двух состояниях на ногу '
  f'(опора/перенос) это 2\u2076{NBSP}={NBSP}64 состояния, а при более '
  f'детальном описании{NBSP}{EMDASH} сотни и тысячи [4].')

P(f'Сети Петри, предложенные К.А.{NBSP}Петри в 1962{NBSP}г. [5], '
  f'позволяют моделировать параллельные процессы без комбинаторного '
  f'взрыва. Математический аппарат сетей Петри (матричное представление, '
  f'S-инварианты, дерево достижимости) даёт возможность формально '
  f'доказывать свойства системы: отсутствие тупиков, ограниченность, '
  f'взаимное исключение [6]. Целью настоящей работы является построение '
  f'формальной модели координации триподной походки гексапода, '
  f'доказательство корректности алгоритма методом S-инвариантов '
  f'и темпоральный анализ для определения предельной скорости перемещения.')


# ══════════════ ТЕОРИЯ ══════════════

SEC('Основные определения',
    f'Сеть Петри определяется как четвёрка [5, 6]:')

OMML('PN = (P, T, F, M\u2080)', 1)

P(f'где P{NBSP}={NBSP}{{p\u2081, p\u2082, \u2026, p\u2098}}{NBSP}'
  f'{EMDASH} конечное множество позиций; '
  f'T{NBSP}={NBSP}{{t\u2081, t\u2082, \u2026, t\u2099}}{NBSP}'
  f'{EMDASH} конечное множество переходов; '
  f'F{NBSP}\u2286{NBSP}(P{NBSP}\u00d7{NBSP}T){NBSP}\u222a{NBSP}'
  f'(T{NBSP}\u00d7{NBSP}P){NBSP}{EMDASH} множество дуг; '
  f'M\u2080:{NBSP}P{NBSP}\u2192{NBSP}\u2115{NBSP}'
  f'{EMDASH} начальная маркировка.')

P(f'Переход t{NBSP}\u2208{NBSP}T разрешён при маркировке M, если '
  f'\u2200p{NBSP}\u2208{NBSP}\u2022t:{NBSP}M(p){NBSP}\u2265{NBSP}1. '
  f'При срабатывании маркировка обновляется:')

OMML("M'(p) = M(p) \u2212 W(p, t) + W(t, p)", 2)

P(f'Матрица инцидентности C{NBSP}={NBSP}C\u207a{NBSP}\u2212{NBSP}C\u207b '
  f'связывает маркировки через вектор срабатываний \u03c3:')

OMML('M = M\u2080 + C \u00b7 \u03c3', 3)

P(f'Вектор x{NBSP}\u2208{NBSP}\u2124\u1d50 называется S-инвариантом, '
  f'если x\u1d40{NBSP}\u00b7{NBSP}C{NBSP}={NBSP}0. Тогда для любой '
  f'достижимой маркировки [7]:')

OMML('x\u1d40 \u00b7 M = x\u1d40 \u00b7 M\u2080 = const', 4)

P(f'Темпоральная сеть Петри дополняет модель функцией задержки '
  f'D:{NBSP}T{NBSP}\u2192{NBSP}\u211d\u207a, сопоставляющей каждому '
  f'переходу время срабатывания [8].')


# ══════════════ МОДЕЛЬ ══════════════

SEC('Описание объекта',
    f'Рассматривается гексапод с шестью идентичными ногами, '
    f'расположенными симметрично по три с каждой стороны корпуса. '
    f'Каждая нога оснащена тремя сервоприводами и может находиться '
    f'в режиме опоры (stance) или переноса (swing). При триподной '
    f'походке ноги разделены на группу{NBSP}A (ноги 1, 3, 5) '
    f'и группу{NBSP}B (ноги 2, 4, 6), работающие в противофазе. '
    f'Критическое ограничение: не менее трёх ног должны быть '
    f'в контакте с поверхностью в любой момент времени [2].')

SEC('Структура сети Петри',
    f'Каждая группа ног проходит три фазы: ожидание (готовность), '
    f'перенос (ноги в воздухе), приземление (установка контакта). '
    f'Координация обеспечивается токеном-мьютексом. '
    f'Позиции модели: p\u2081{NBSP}{EMDASH} готовность{NBSP}A, '
    f'p\u2082{NBSP}{EMDASH} перенос{NBSP}A, '
    f'p\u2083{NBSP}{EMDASH} приземление{NBSP}A, '
    f'p\u2084{NBSP}{EMDASH} готовность{NBSP}B, '
    f'p\u2085{NBSP}{EMDASH} перенос{NBSP}B, '
    f'p\u2086{NBSP}{EMDASH} приземление{NBSP}B, '
    f'p\u2087{NBSP}{EMDASH} мьютекс {LAQUO}устойчивость{RAQUO}.')

P(f'Переходы: '
  f't\u2081 ({LAQUO}отрыв{NBSP}A{RAQUO}): '
  f'p\u2081{NBSP}+{NBSP}p\u2087{NBSP}\u2192{NBSP}p\u2082; '
  f't\u2082 ({LAQUO}перенос{NBSP}A{RAQUO}): '
  f'p\u2082{NBSP}\u2192{NBSP}p\u2083; '
  f't\u2083 ({LAQUO}контакт{NBSP}A{RAQUO}): '
  f'p\u2083{NBSP}\u2192{NBSP}p\u2081{NBSP}+{NBSP}p\u2087; '
  f't\u2084 ({LAQUO}отрыв{NBSP}B{RAQUO}): '
  f'p\u2084{NBSP}+{NBSP}p\u2087{NBSP}\u2192{NBSP}p\u2085; '
  f't\u2085 ({LAQUO}перенос{NBSP}B{RAQUO}): '
  f'p\u2085{NBSP}\u2192{NBSP}p\u2086; '
  f't\u2086 ({LAQUO}контакт{NBSP}B{RAQUO}): '
  f'p\u2086{NBSP}\u2192{NBSP}p\u2084{NBSP}+{NBSP}p\u2087.')

P(f'Начальная маркировка: M\u2080{NBSP}={NBSP}(1, 0, 0, 1, 0, 0, 1). '
  f'Из начального состояния разрешены t\u2081 и t\u2084, но при '
  f'срабатывании одного из них токен p\u2087 изымается и второй '
  f'переход блокируется до его возврата.')

P(f'[Рис.{NBSP}1. Сеть Петри координации триподной походки гексапода]',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=6, after=6)


# ══════════════ АНАЛИЗ ══════════════

SEC('Матрица инцидентности',
    f'Матрица C размером 7{NBSP}\u00d7{NBSP}6 построена по формуле (3):')

make_table(
    headers=['', 't\u2081', 't\u2082', 't\u2083', 't\u2084', 't\u2085', 't\u2086'],
    rows=[
        ['p\u2081', '\u22121', '0', '1', '0', '0', '0'],
        ['p\u2082', '1', '\u22121', '0', '0', '0', '0'],
        ['p\u2083', '0', '1', '\u22121', '0', '0', '0'],
        ['p\u2084', '0', '0', '0', '\u22121', '0', '1'],
        ['p\u2085', '0', '0', '0', '1', '\u22121', '0'],
        ['p\u2086', '0', '0', '0', '0', '1', '\u22121'],
        ['p\u2087', '\u22121', '0', '1', '\u22121', '0', '1'],
    ],
)

P('', after=4)

SEC('S-инварианты',
    f'Решая систему x\u1d40{NBSP}\u00b7{NBSP}C{NBSP}={NBSP}0, '
    f'получаем три линейно независимых инварианта:')

OMML('x\u2081 = (1, 1, 1, 0, 0, 0, 0)\u1d40', 5)
OMML('x\u2082 = (0, 0, 0, 1, 1, 1, 0)\u1d40', 6)
OMML('x\u2083 = (0, 1, 1, 0, 1, 1, 1)\u1d40', 7)

P(f'Носители инвариантов покрывают все позиции сети{NBSP}{EMDASH} '
  f'сеть консервативна [6]. Инвариант x\u2081 даёт:')

OMML('M(p\u2081) + M(p\u2082) + M(p\u2083) = 1', 8)

P(f'Группа{NBSP}A всегда ровно в одной фазе. Аналогично x\u2082 для '
  f'группы{NBSP}B. Инвариант x\u2083{NBSP}{EMDASH} ключевой:')

OMML('M(p\u2082) + M(p\u2083) + M(p\u2085) + M(p\u2086) + M(p\u2087) = 1', 9)

P(f'Поскольку маркировка неотрицательна, отсюда M(p\u2082){NBSP}+{NBSP}'
  f'M(p\u2085){NBSP}\u2264{NBSP}1. Это формально доказывает взаимное '
  f'исключение: обе группы не могут одновременно находиться в переносе. '
  f'В физическом смысле{NBSP}{EMDASH} не менее трёх ног всегда '
  f'в контакте с поверхностью, что гарантирует статическую устойчивость. '
  f'Доказательство получено без перебора маркировок{NBSP}{EMDASH} оно '
  f'следует из структуры сети [7].')

SEC('Отсутствие тупиков',
    f'Покажем живость сети. Каждый цикл '
    f'(t\u2081\u2192t\u2082\u2192t\u2083 или '
    f't\u2084\u2192t\u2085\u2192t\u2086) гарантированно завершается '
    f'и возвращает токен мьютекса в p\u2087. Из любой достижимой '
    f'маркировки хотя бы один переход разрешён. Сеть является живой '
    f'и 1-ограниченной (безопасной).')

SEC('Трассировка маркировок',
    f'Для наглядности приведём полный цикл (табл.{NBSP}1).')

make_table(
    headers=['Шаг', 'Переход', 'p\u2081', 'p\u2082', 'p\u2083',
             'p\u2084', 'p\u2085', 'p\u2086', 'p\u2087', 'Комментарий'],
    rows=[
        ['0', EMDASH, '1', '0', '0', '1', '0', '0', '1', 'Начальное состояние'],
        ['1', 't\u2081', '0', '1', '0', '1', '0', '0', '0', 'A оторвалась'],
        ['2', 't\u2082', '0', '0', '1', '1', '0', '0', '0', 'A перенесена'],
        ['3', 't\u2083', '1', '0', '0', '1', '0', '0', '1', 'A приземлилась'],
        ['4', 't\u2084', '1', '0', '0', '0', '1', '0', '0', 'B оторвалась'],
        ['5', 't\u2085', '1', '0', '0', '0', '0', '1', '0', 'B перенесена'],
        ['6', 't\u2086', '1', '0', '0', '1', '0', '0', '1', f'B приземлилась = M\u2080'],
    ],
    caption=f'Табл.{NBSP}1. Трассировка маркировок при полном цикле походки'
)

P('', after=4)

P(f'После шага 6 система возвращается в M\u2080. На протяжении всей '
  f'трассировки M(p\u2082){NBSP}+{NBSP}M(p\u2085){NBSP}\u2264{NBSP}1, '
  f'что подтверждает аналитический результат.')


# ══════════════ ТЕМПОРАЛЬНЫЙ АНАЛИЗ ══════════════

SEC('Темпоральный анализ',
    f'Присвоим переходам задержки, соответствующие реальным временам '
    f'фаз походки гексапода массой 2{NBSP}кг с сервоприводами класса '
    f'25{NBSP}кг\u00b7см [8]: '
    f'd(t\u2081){NBSP}={NBSP}d(t\u2084){NBSP}={NBSP}50{NBSP}мс (отрыв), '
    f'd(t\u2082){NBSP}={NBSP}d(t\u2085){NBSP}={NBSP}200{NBSP}мс (перенос), '
    f'd(t\u2083){NBSP}={NBSP}d(t\u2086){NBSP}={NBSP}50{NBSP}мс (приземление).')

P(f'Полный цикл группы: 50{NBSP}+{NBSP}200{NBSP}+{NBSP}50{NBSP}={NBSP}'
  f'300{NBSP}мс. Поскольку группы работают последовательно, полный '
  f'период походки:')

OMML('T = T_A + T_B = 300 + 300 = 600 \u043c\u0441', 10)

P(f'За один период робот совершает два шага длиной '
  f'l{NBSP}={NBSP}60{NBSP}мм. Скорость:')

OMML('V = 2l / T = 2 \u00d7 60 / 600 = 200 \u043c\u043c/\u0441', 11)

P(f'Это согласуется с типичными значениями 150{EMDASH}300{NBSP}мм/с '
  f'для гексаподов данного класса [9].')

SEC('Узкое место',
    f'Фаза переноса (t\u2082, t\u2085) занимает 200{NBSP}мс из '
    f'300{NBSP}мс полуцикла{NBSP}{EMDASH} 67{NBSP}%. '
    f'Для увеличения скорости в первую очередь следует сокращать '
    f'именно время переноса.')

SEC('Оценка модернизации',
    f'Замена сервоприводов с сокращением времени переноса на 25{NBSP}% '
    f'(d(t\u2082){NBSP}={NBSP}d(t\u2085){NBSP}={NBSP}150{NBSP}мс):')

OMML("T' = 2 \u00d7 (50 + 150 + 50) = 500 \u043c\u0441", 12)

OMML("V' = 120 / 0.5 = 240 \u043c\u043c/\u0441", 13)

P(f'Прирост скорости{NBSP}{EMDASH} 20{NBSP}%. Итоговое ускорение '
  f'меньше 25{NBSP}%, поскольку фиксированные фазы отрыва и '
  f'приземления {LAQUO}разбавляют{RAQUO} эффект. Такие оценки позволяют '
  f'обоснованно выбирать компоненты на этапе проектирования.')


# ══════════════ ВЫВОДЫ ══════════════

SEC('Выводы',
    f'Построена формальная модель координации триподной походки '
    f'гексапода (7{NBSP}позиций, 6{NBSP}переходов). Методом S-инвариантов '
    f'доказано взаимное исключение фаз переноса, гарантирующее '
    f'статическую устойчивость. Показана живость и 1-ограниченность сети. '
    f'Темпоральный анализ определил период походки (600{NBSP}мс), '
    f'скорость (200{NBSP}мм/с) и узкое место (перенос{NBSP}{EMDASH} '
    f'67{NBSP}%). Замена сервоприводов даёт +20{NBSP}% скорости. '
    f'Направлением дальнейшей работы является применение стохастических '
    f'сетей Петри для учёта вариативности времени срабатывания '
    f'приводов [10].')


# ══════════════ ЛИТЕРАТУРА ══════════════

P('Литература', bold=True, indent=False, before=12, after=6)

refs_ru = [
    f'Siciliano{NBSP}B., Khatib{NBSP}O. (eds.) Springer Handbook of Robotics. '
    f'Springer, 2016, 2228{NBSP}p.',

    f'Belter{NBSP}D., Skrzypczy\u0144ski{NBSP}P. A biologically inspired approach '
    f'to feasible gait learning for a hexapod robot. Applied Mathematics and '
    f'Computer Science, 2010, vol.{NBSP}20, no.{NBSP}1, pp.{NBSP}69{EMDASH}84.',

    f'Lee{NBSP}E.A., Seshia{NBSP}S.A. Introduction to Embedded Systems: '
    f'A Cyber-Physical Systems Approach. MIT Press, 2017, 568{NBSP}p.',

    f'Hopcroft{NBSP}J.E., Motwani{NBSP}R., Ullman{NBSP}J.D. '
    f'Introduction to Automata Theory, Languages, and Computation. '
    f'Pearson, 2006, 535{NBSP}p.',

    f'Petri{NBSP}C.A. Kommunikation mit Automaten: Dissertation. '
    f'Universit\u00e4t Bonn, 1962, 128{NBSP}S.',

    f'Murata{NBSP}T. Petri Nets: Properties, Analysis and Applications. '
    f'Proceedings of the IEEE, 1989, vol.{NBSP}77, no.{NBSP}4, pp.{NBSP}541{EMDASH}580.',

    f'Silva{NBSP}M., Teruel{NBSP}E., Colom{NBSP}J.M. Linear Algebraic and '
    f'Linear Programming Techniques for the Analysis of Place/Transition Net Systems. '
    f'Lectures on Petri Nets I: Basic Models. Springer, 1998, pp.{NBSP}309{EMDASH}373.',

    f'Ramchandani{NBSP}C. Analysis of Asynchronous Concurrent Systems by Timed '
    f'Petri Nets. Technical Report MAC-TR-120. MIT, 1974, 106{NBSP}p.',

    f'Спирин{NBSP}Н.А., Лавров{NBSP}В.В. Управление движением '
    f'многоногих шагающих роботов. Мехатроника, автоматизация, '
    f'управление, 2019, т.{NBSP}20, \u2116{NBSP}7, с.{NBSP}430{EMDASH}438.',

    f'Marsan{NBSP}M.A., Balbo{NBSP}G., Conte{NBSP}G. et{NBSP}al. '
    f'Modelling with Generalized Stochastic Petri Nets. '
    f'John Wiley & Sons, 1995, 316{NBSP}p.',
]

for i, ref in enumerate(refs_ru, 1):
    P(f'[{i}] {ref}', indent=False)

# ── Сведения об авторе ──
P('', after=8)

P(f'Романов Владимир Павлович{NBSP}{EMDASH} студент кафедры '
  f'{LAQUO}Прикладная робототехника{RAQUO}, '
  f'МГТУ им.{NBSP}Н.Э.{NBSP}Баумана, Москва, Российская Федерация.',
  indent=False)

P('', after=4)

P(f'Научный руководитель{NBSP}{EMDASH} Шевченко Максим Юрьевич, '
  f'кафедра {LAQUO}Прикладная робототехника{RAQUO}, '
  f'МГТУ им.{NBSP}Н.Э.{NBSP}Баумана, Москва, Российская Федерация.',
  indent=False)

P('', after=4)

# Ссылка на цитирование
RP([
    ('Ссылку на эту статью просим оформлять следующим образом:', True, False),
], indent=False, after=2)

P(f'Романов{NBSP}В.П. Формальная верификация координации походки '
  f'шестиногого робота методом сетей Петри. Политехнический молодёжный '
  f'журнал, 2026.', indent=False)


# ══════════════════════════════════════════════════════════
#  АНГЛИЙСКИЙ БЛОК (новая страница)
# ══════════════════════════════════════════════════════════

doc.add_page_break()

P('FORMAL VERIFICATION OF HEXAPOD ROBOT\n'
  'GAIT COORDINATION USING PETRI NETS',
  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=4)

P(f'V.P.{NBSP}Romanov', indent=False, after=0)
P('romanov.vp@student.bmstu.ru', indent=False, after=2, sz=Pt(11))
P('Bauman Moscow State Technical University, Moscow, Russian Federation',
  indent=False, after=8)

abstract_en = (
    'The paper addresses formal verification of the tripod gait coordination '
    'algorithm for a hexapod walking robot using Petri nets. The robot\'s six '
    'legs are divided into two groups of three, alternating between swing and '
    'stance phases. A Petri net model is constructed that enforces static '
    'stability: at least three legs must maintain ground contact at all times. '
    'Using S-invariants, mutual exclusion of swing phases is formally proven '
    'and the absence of deadlocks is demonstrated. Temporal analysis with '
    'deterministic delays determines the gait cycle time (600 ms), maximum '
    'locomotion speed (200 mm/s), and the bottleneck (swing phase accounts '
    'for 67% of cycle time). It is shown that replacing servos to reduce '
    'swing time by 25% yields a 20% speed improvement.'
)

keywords_en = (
    'Petri nets, walking robot, hexapod, tripod gait, formal verification, '
    'S-invariants, gait coordination, static stability, temporal analysis, '
    'mutual exclusion'
)

abstract_block_en(abstract_en, keywords_en,
                  '21.05.2026',
                  f'\u00a9 Bauman Moscow State Technical University, 2026')

P('', after=6)

# ── English references ──
P('References', bold=True, indent=False, before=12, after=6)

refs_en = [
    f'Siciliano{NBSP}B., Khatib{NBSP}O. (eds.) Springer Handbook of Robotics. '
    f'Springer, 2016, 2228{NBSP}p.',

    f'Belter{NBSP}D., Skrzypczy\u0144ski{NBSP}P. A biologically inspired approach '
    f'to feasible gait learning for a hexapod robot. Applied Mathematics and '
    f'Computer Science, 2010, vol.{NBSP}20, no.{NBSP}1, pp.{NBSP}69{EMDASH}84.',

    f'Lee{NBSP}E.A., Seshia{NBSP}S.A. Introduction to Embedded Systems: '
    f'A Cyber-Physical Systems Approach. MIT Press, 2017, 568{NBSP}p.',

    f'Hopcroft{NBSP}J.E., Motwani{NBSP}R., Ullman{NBSP}J.D. '
    f'Introduction to Automata Theory, Languages, and Computation. '
    f'Pearson, 2006, 535{NBSP}p.',

    f'Petri{NBSP}C.A. Kommunikation mit Automaten: Dissertation. '
    f'Universit\u00e4t Bonn, 1962, 128{NBSP}S.',

    f'Murata{NBSP}T. Petri Nets: Properties, Analysis and Applications. '
    f'Proceedings of the IEEE, 1989, vol.{NBSP}77, no.{NBSP}4, pp.{NBSP}541{EMDASH}580.',

    f'Silva{NBSP}M., Teruel{NBSP}E., Colom{NBSP}J.M. Linear Algebraic and '
    f'Linear Programming Techniques for the Analysis of Place/Transition Net Systems. '
    f'Lectures on Petri Nets I: Basic Models. Springer, 1998, pp.{NBSP}309{EMDASH}373.',

    f'Ramchandani{NBSP}C. Analysis of Asynchronous Concurrent Systems by Timed '
    f'Petri Nets. Technical Report MAC-TR-120. MIT, 1974, 106{NBSP}p.',

    f'Spirin{NBSP}N.A., Lavrov{NBSP}V.V. Upravlenie dvizheniem '
    f'shagayushchikh robotov [Motion control of walking robots]. '
    f'Mekhatronika, avtomatizatsiya, upravlenie [Mechatronics, '
    f'automation, control], 2019, vol.{NBSP}20, no.{NBSP}7, '
    f'pp.{NBSP}430{EMDASH}438 (in Russ.).',

    f'Marsan{NBSP}M.A., Balbo{NBSP}G., Conte{NBSP}G. et{NBSP}al. '
    f'Modelling with Generalized Stochastic Petri Nets. '
    f'John Wiley & Sons, 1995, 316{NBSP}p.',
]

for i, ref in enumerate(refs_en, 1):
    P(f'[{i}] {ref}', indent=False)

# ── Author info (English) ──
P('', after=8)

P(f'Romanov V.P.{NBSP}{EMDASH} Student, Department of Applied Robotics, '
  f'Bauman Moscow State Technical University, Moscow, Russian Federation.',
  indent=False)

P('', after=4)

P(f'Scientific advisor{NBSP}{EMDASH} Shevchenko M.Yu., Department of '
  f'Applied Robotics, Bauman Moscow State Technical University, Moscow, '
  f'Russian Federation.',
  indent=False)

P('', after=4)

RP([
    ('Please cite this article in English as:', True, False),
], indent=False, after=2)

P(f'Romanov{NBSP}V.P. Formal verification of hexapod robot gait coordination '
  f'using Petri nets. Politekhnicheskiy molodezhnyy zhurnal '
  f'[Politechnical student journal], 2026.', indent=False)


# ── Нумерация ──
add_page_numbers()

# ── Сохранение ──
out = os.path.join(os.path.dirname(__file__), 'Романов.docx')
doc.save(out)
print(f'Статья сохранена: {out}')
